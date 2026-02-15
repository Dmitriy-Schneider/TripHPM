"""
Упрощенный генератор документов - работает только с готовыми шаблонами
Простая логика: копирование шаблона + заполнение ячеек + архивирование

РАЗДЕЛЕНИЕ НА ДВА ЭТАПА:
1. ДО поездки: Приказ + СЗ (запрос аванса)
2. ПОСЛЕ поездки: АО + СЗ (только при перерасходе)
"""
from pathlib import Path
from datetime import date, datetime, timedelta
from typing import Dict, Iterable, Optional
from decimal import Decimal, ROUND_CEILING
import re
import zipfile
import shutil

from docx import Document
from openpyxl import load_workbook
from num2words import num2words
import logging

# win32com для работы с Excel (не портит файлы в отличие от openpyxl/xlwings)
try:
    import win32com.client as win32
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False


class SimpleDocumentGenerator:
    """Упрощенный генератор документов с разделением на этапы"""

    def __init__(self, templates_dir: Path, output_dir: Path):
        self.templates_dir = Path(templates_dir)
        self.output_dir = Path(output_dir)

    def _get_trip_folder(self, trip_data: Dict) -> Path:
        """Возвращает путь к папке командировки"""
        folder_name = f"{trip_data['date_from'].strftime('%Y-%m-%d')}_{trip_data['destination_city']}"
        return self.output_dir / folder_name

    def _ensure_folder_structure(self, trip_folder: Path) -> None:
        """Создает структуру папок если нужно"""
        trip_folder.mkdir(parents=True, exist_ok=True)
        (trip_folder / "documents").mkdir(exist_ok=True)
        (trip_folder / "receipts").mkdir(exist_ok=True)

    # ==================== ЭТАП 1: ДО ПОЕЗДКИ ====================

    def generate_pre_trip(self, trip_data: Dict, custom_output_dir: Optional[Path] = None) -> Dict[str, str]:
        """
        Генерирует документы ДО поездки:
        - Приказ
        - Служебная записка (запрос аванса)

        Дата документов = prikaz_date или sz_date из trip_data, или сегодня
        Сумма в СЗ = advance_rub (запрашиваемый аванс)

        БЕЗ ZIP архива - только отдельные файлы для скачивания
        """
        # Используем пользовательскую директорию если указана
        if custom_output_dir:
            trip_folder = Path(custom_output_dir) / self._get_folder_name(trip_data)
        else:
            trip_folder = self._get_trip_folder(trip_data)

        self._ensure_folder_structure(trip_folder)

        results = {
            'folder': str(trip_folder)
        }

        # 1. Приказ (дата = prikaz_date или сегодня)
        results['prikaz'] = self._generate_prikaz(trip_data, trip_folder)

        # 2. Служебная записка на аванс
        results['sz_advance'] = self._generate_sz_advance(trip_data, trip_folder)

        # Этап 1 НЕ создает ZIP - только файлы
        return results

    def _get_folder_name(self, trip_data: Dict) -> str:
        """Возвращает имя папки для командировки"""
        return f"{trip_data['date_from'].strftime('%Y-%m-%d')}_{trip_data['destination_city']}"

    def _generate_sz_advance(self, trip_data: Dict, output_folder: Path) -> str:
        """
        Генерирует Служебную записку на запрос аванса ДО поездки.
        Сумма = advance_rub (только аванс, без чеков).
        Дата = sz_date или сегодня.
        """
        template_path = self.templates_dir / "sz_template.docx"
        output_path = output_folder / "documents" / "Служебная_записка_аванс.docx"

        logger = logging.getLogger(__name__)
        doc = Document(template_path)

        org_name = trip_data.get('org_name', 'ООО «ВЭМ»')
        fio = trip_data['fio']
        fio_short = self._fio_short(fio)

        # Дата = sz_date или сегодня
        doc_date = trip_data.get('sz_date') or date.today()
        if isinstance(doc_date, datetime):
            doc_date = doc_date.date()

        # Сумма = только аванс
        advance_amount = trip_data.get('advance_rub', 0) or 0
        advance_rounded = self._ceil_rubles(advance_amount)
        advance_words = num2words(advance_rounded, lang='ru')

        logger.info(
            "[SZ_ADVANCE] advance=%s rounded=%s words='%s' doc_date=%s",
            advance_amount, advance_rounded, advance_words, doc_date.strftime('%d.%m.%Y')
        )

        sum_pattern = re.compile(r'сумму\s+\d+[\d\s]*\s*\([^)]*\)\s+рублей', re.IGNORECASE)
        date_pattern = re.compile(r'\d{2}\.\d{2}\.\d{4}')

        for paragraph in doc.paragraphs:
            raw_text = paragraph.text

            # НЕ меняем ФИО - берем из шаблона как есть!
            # Меняем только сумму и дату

            updated = sum_pattern.sub(
                f"сумму {advance_rounded} ({advance_words}) рублей",
                raw_text
            )
            updated = date_pattern.sub(doc_date.strftime('%d.%m.%Y'), updated)

            if updated != raw_text:
                paragraph.text = updated

        doc.save(output_path)
        return str(output_path)

    # ==================== ЭТАП 2: ПОСЛЕ ПОЕЗДКИ ====================

    def generate_post_trip(self, trip_data: Dict, custom_output_dir: Optional[Path] = None) -> Dict[str, str]:
        """
        Генерирует документы ПОСЛЕ поездки:
        - Авансовый отчет
        - Служебная записка на доплату (ТОЛЬКО если перерасход > 0)

        Возвращает словарь с путями к файлам и флагом needs_sz_dopay
        """
        # Используем пользовательскую директорию если указана
        if custom_output_dir:
            trip_folder = Path(custom_output_dir) / self._get_folder_name(trip_data)
        else:
            trip_folder = self._get_trip_folder(trip_data)

        self._ensure_folder_structure(trip_folder)

        results = {
            'folder': str(trip_folder)
        }

        # 1. Авансовый отчет
        results['ao'] = self._generate_ao(trip_data, trip_folder)

        # 2. Проверяем: нужна ли СЗ на доплату?
        # to_return > 0 = остаток к возврату (не нужна СЗ)
        # to_return < 0 = перерасход, нужна доплата (нужна СЗ)
        to_return = trip_data.get('to_return', 0)
        results['to_return'] = to_return
        results['needs_sz_dopay'] = to_return < 0

        if to_return < 0:
            # Перерасход - генерируем СЗ на доплату
            results['sz_dopay'] = self._generate_sz_dopay(trip_data, trip_folder, abs(to_return))
        else:
            results['sz_dopay'] = None

        # 3. Копируем чеки
        self._copy_receipts(trip_data, trip_folder)

        # 4. Создаем ZIP
        results['zip'] = self._create_zip(trip_folder)

        return results

    def _generate_sz_dopay(self, trip_data: Dict, output_folder: Path, dopay_amount: float) -> str:
        """
        Генерирует Служебную записку на доплату ПОСЛЕ поездки.
        Создается только если перерасход > 0.
        Сумма = сумма доплаты (перерасход).
        Дата = ao_date или сегодня.
        """
        template_path = self.templates_dir / "sz_template.docx"
        output_path = output_folder / "documents" / "Служебная_записка_доплата.docx"

        logger = logging.getLogger(__name__)
        doc = Document(template_path)

        org_name = trip_data.get('org_name', 'ООО «ВЭМ»')
        fio = trip_data['fio']
        fio_short = self._fio_short(fio)

        # Дата = ao_date или сегодня
        doc_date = trip_data.get('ao_date') or date.today()
        if isinstance(doc_date, datetime):
            doc_date = doc_date.date()

        dopay_rounded = self._ceil_rubles(dopay_amount)
        dopay_words = num2words(dopay_rounded, lang='ru')

        logger.info(
            "[SZ_DOPAY] dopay=%s rounded=%s words='%s' doc_date=%s",
            dopay_amount, dopay_rounded, dopay_words, doc_date.strftime('%d.%m.%Y')
        )

        sum_pattern = re.compile(r'сумму\s+\d+[\d\s]*\s*\([^)]*\)\s+рублей', re.IGNORECASE)
        date_pattern = re.compile(r'\d{2}\.\d{2}\.\d{4}')

        for paragraph in doc.paragraphs:
            raw_text = paragraph.text

            # НЕ меняем ФИО - берем из шаблона как есть!
            # Меняем только сумму и дату

            updated = sum_pattern.sub(
                f"сумму {dopay_rounded} ({dopay_words}) рублей",
                raw_text
            )
            updated = date_pattern.sub(doc_date.strftime('%d.%m.%Y'), updated)

            if updated != raw_text:
                paragraph.text = updated

        doc.save(output_path)
        return str(output_path)

    # ==================== СТАРЫЙ МЕТОД (для обратной совместимости) ====================

    def generate_all(self, trip_data: Dict) -> Dict[str, str]:
        """
        Генерирует все документы - для обратной совместимости.
        Рекомендуется использовать generate_pre_trip и generate_post_trip.
        """
        trip_folder = self._get_trip_folder(trip_data)

        # Удаляем старую папку если есть
        if trip_folder.exists():
            shutil.rmtree(trip_folder)

        self._ensure_folder_structure(trip_folder)

        results = {}

        # 1. Приказ
        results['prikaz'] = self._generate_prikaz(trip_data, trip_folder)

        # 2. Авансовый отчет
        results['ao'] = self._generate_ao(trip_data, trip_folder)

        # 3. Служебная записка (старый вариант с total_expenses)
        results['sz'] = self._generate_sz(trip_data, trip_folder)

        # 4. Копируем чеки
        self._copy_receipts(trip_data, trip_folder)

        # 5. Создаем ZIP
        results['zip'] = self._create_zip(trip_folder)

        return results

    def _generate_prikaz(self, trip_data: Dict, output_folder: Path) -> str:
        """Генерирует Приказ - заполняет только даты, место и цель командировки.
        ФИО, табельный номер, отдел, должность - берутся из шаблона без изменений!
        """
        logger = logging.getLogger(__name__)
        template_path = self.templates_dir / "prikaz_template.docx"
        output_path = output_folder / "documents" / "Приказ.docx"

        doc = Document(template_path)

        destination_org = trip_data['destination_org']
        purpose = trip_data['purpose']
        days = str(trip_data['days'])

        # Конвертируем даты из строк если нужно
        date_from = trip_data['date_from']
        date_to = trip_data['date_to']

        logger.info("[PRIKAZ] Raw date_from=%s (type=%s), date_to=%s (type=%s)",
                    date_from, type(date_from).__name__, date_to, type(date_to).__name__)

        if isinstance(date_from, str):
            date_from = datetime.strptime(date_from, '%Y-%m-%d').date()
        elif isinstance(date_from, datetime):
            date_from = date_from.date()

        if isinstance(date_to, str):
            date_to = datetime.strptime(date_to, '%Y-%m-%d').date()
        elif isinstance(date_to, datetime):
            date_to = date_to.date()

        logger.info("[PRIKAZ] Converted date_from=%s, date_to=%s, days=%s",
                    date_from, date_to, days)

        # Дата приказа = prikaz_date или сегодня
        order_date = trip_data.get('prikaz_date') or date.today()
        if isinstance(order_date, str):
            order_date = datetime.strptime(order_date, '%Y-%m-%d').date()
        elif isinstance(order_date, datetime):
            order_date = order_date.date()

        tables = doc.tables
        logger.info("[PRIKAZ] Found %d tables in document", len(tables))
        if len(tables) >= 8:
            # НЕ меняем: организация, ФИО, табельный - берем из шаблона!
            # Дата составления
            tables[1].cell(1, 2).text = order_date.strftime('%d.%m.%Y')
            # Место назначения
            tables[3].cell(0, 0).text = destination_org
            # Количество дней
            tables[4].cell(0, 1).text = days
            # Даты командировки
            logger.info("[PRIKAZ] Filling dates in table[5]: %s - %s", date_from, date_to)
            self._fill_split_date(tables[5], date_from, date_to)
            # Цель
            tables[6].cell(0, 1).text = purpose
            # Дата ознакомления работника
            if len(tables) >= 10:
                self._fill_ack_date(tables[9], order_date)

        # НЕ меняем отдел и должность - берем из шаблона!

        doc.save(output_path)
        return str(output_path)

    def _generate_ao(self, trip_data: Dict, output_folder: Path) -> str:
        """Генерирует Авансовый отчет через win32com (не портит файлы)"""
        template_path = self.templates_dir / "ao_template.xlsx"
        output_path = output_folder / "documents" / "Авансовый_отчет.xlsx"

        logger = logging.getLogger(__name__)

        if not HAS_WIN32COM:
            logger.error("[AO] win32com not available!")
            raise RuntimeError("win32com не установлен. Установите: pip install pywin32")

        logger.info("[AO] Generating with win32com...")

        # Копируем шаблон
        shutil.copy(template_path, output_path)

        # Открываем через win32com
        excel = win32.Dispatch("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False

        try:
            wb = excel.Workbooks.Open(str(output_path.resolve()))
            ws = wb.Sheets(1)

            # Дата оформления документа (ячейка Z13) = ao_date или сегодня
            doc_date = trip_data.get('ao_date')
            if doc_date:
                if isinstance(doc_date, str):
                    doc_date = datetime.strptime(doc_date, '%Y-%m-%d').date()
                elif isinstance(doc_date, datetime):
                    doc_date = doc_date.date()
            else:
                doc_date = date.today()

            ws.Range('Z13').Value = doc_date.strftime('%d.%m.%y')

            # Маппинг категорий
            category_names = {
                'taxi': 'Такси',
                'fuel': 'Топливо',
                'hotel': 'Гостиница',
                'restaurant': 'Автобус',
                'bus': 'Автобус',
                'flight': 'Самолет',
                'airplane': 'Самолет',
                'train': 'Поезд',
                'самолет': 'Самолет',
                'поезд': 'Поезд',
                'автобус': 'Автобус',
                'other': 'Представительские'
            }
            category_order = ['fuel', 'taxi', 'flight', 'airplane', 'train', 'bus', 'hotel', 'restaurant', 'other']

            # Очищаем диапазон расходов
            for clear_row in range(63, 85):
                ws.Range(f'P{clear_row}').Value = None
                ws.Range(f'Y{clear_row}').Value = None

            # Заполняем расходы
            row = 63
            raw_expenses = trip_data.get('expenses_by_category', {})
            expenses = {}
            for key, amount in raw_expenses.items():
                normalized_key = self._normalize_category_key(key)
                expenses[normalized_key] = (expenses.get(normalized_key, 0) or 0) + (amount or 0)

            for category in self._ordered_categories(expenses.keys(), category_order):
                amount = expenses.get(category)
                if amount and amount > 0:
                    normalized_category = self._normalize_category_key(category)
                    display_category = category_names.get(normalized_category, category)
                    ws.Range(f'P{row}').Value = display_category
                    ws.Range(f'Y{row}').Value = float(self._to_money(amount))
                    row += 1

            # Суточные
            per_diem = trip_data.get('per_diem_to_pay', 0)
            if per_diem > 0:
                ws.Range(f'P{row}').Value = 'Суточные'
                ws.Range(f'Y{row}').Value = float(self._to_money(per_diem))

            wb.Save()
            wb.Close()
            logger.info("[AO] File saved successfully: %s", output_path)
            return str(output_path)
        finally:
            excel.Quit()

    def _generate_sz(self, trip_data: Dict, output_folder: Path) -> str:
        """Генерирует Служебную записку (старый метод для совместимости)
        ФИО берется из шаблона - НЕ заменяется!
        """
        template_path = self.templates_dir / "sz_template.docx"
        output_path = output_folder / "documents" / "Служебная_записка.docx"

        logger = logging.getLogger(__name__)
        doc = Document(template_path)

        doc_date = self._calculate_doc_date(trip_data['date_from'])

        total_expenses = trip_data.get('total_expenses', 0)
        total_expenses_rounded = self._ceil_rubles(total_expenses)
        total_expenses_words = num2words(total_expenses_rounded, lang='ru')
        logger.info(
            "[SZ] total_expenses=%s rounded=%s words='%s' doc_date=%s template=%s output=%s",
            total_expenses,
            total_expenses_rounded,
            total_expenses_words,
            doc_date.strftime('%d.%m.%Y'),
            template_path,
            output_path
        )

        sum_pattern = re.compile(r'сумму\s+\d+[\d\s]*\s*\([^)]*\)\s+рублей', re.IGNORECASE)
        date_pattern = re.compile(r'\d{2}\.\d{2}\.\d{4}')

        # НЕ меняем ФИО - берем из шаблона как есть!
        # Меняем только сумму и дату
        for paragraph in doc.paragraphs:
            raw_text = paragraph.text

            updated = sum_pattern.sub(
                f"сумму {total_expenses_rounded} ({total_expenses_words}) рублей",
                raw_text
            )
            updated = date_pattern.sub(doc_date.strftime('%d.%m.%Y'), updated)

            if updated != raw_text:
                paragraph.text = updated

        doc.save(output_path)
        return str(output_path)

    @staticmethod
    def _ordered_categories(categories: Iterable[str], order: Iterable[str]) -> Iterable[str]:
        seen = set()
        ordered = []
        for key in order:
            if key in categories:
                ordered.append(key)
                seen.add(key)
        for key in categories:
            if key not in seen:
                ordered.append(key)
        return ordered

    @staticmethod
    def _normalize_category_key(category: str) -> str:
        if category is None:
            return 'other'
        raw = str(category).strip()
        lower = raw.lower()
        mapping = {
            'самолет': 'airplane',
            'airplane': 'airplane',
            'flight': 'airplane',
            'поезд': 'train',
            'train': 'train',
            'такси': 'taxi',
            'taxi': 'taxi',
            'топливо': 'fuel',
            'fuel': 'fuel',
            'гостиница': 'hotel',
            'hotel': 'hotel',
            'автобус': 'bus',
            'bus': 'bus',
            'ресторан': 'bus',
            'restaurant': 'bus',
            'представительские': 'other',
            'other': 'other',
        }
        return mapping.get(lower, raw)

    @staticmethod
    def _to_money(value: float) -> float:
        return float(Decimal(str(value)).quantize(Decimal("0.01")))

    @staticmethod
    def _ceil_rubles(value: float) -> int:
        return int(Decimal(str(value)).quantize(Decimal("1"), rounding=ROUND_CEILING))

    @staticmethod
    def _calculate_doc_date(date_from: date) -> date:
        doc_date = date_from - timedelta(days=2)
        while doc_date.weekday() >= 5:
            doc_date -= timedelta(days=1)
        return doc_date

    @staticmethod
    def _fio_short(fio: str) -> str:
        parts = fio.strip().split()
        if not parts:
            return fio
        surname = parts[0]
        initials = "".join([f"{p[0]}." for p in parts[1:3] if p])
        return f"{surname} {initials}".strip()

    @staticmethod
    def _fill_split_date(table, date_from: date, date_to: date):
        year_from = date_from.strftime('%Y')
        year_to = date_to.strftime('%Y')
        table.cell(0, 2).text = date_from.strftime('%d')
        table.cell(0, 4).text = date_from.strftime('%m')
        table.cell(0, 5).text = year_from[:2]
        table.cell(0, 6).text = year_from[2:]
        table.cell(0, 8).text = date_to.strftime('%d')
        table.cell(0, 10).text = date_to.strftime('%m')
        table.cell(0, 11).text = year_to[:2]
        table.cell(0, 12).text = year_to[2:]

    @staticmethod
    def _fill_ack_date(table, ack_date: date):
        year = ack_date.strftime('%Y')
        table.cell(0, 3).text = ack_date.strftime('%d')
        table.cell(0, 5).text = ack_date.strftime('%m')
        table.cell(0, 6).text = year[:2]
        table.cell(0, 7).text = year[2:]

    def _copy_receipts(self, trip_data: Dict, output_folder: Path):
        """Копирует чеки"""
        receipts_folder = output_folder / "receipts"
        receipts = trip_data.get('receipts', [])

        for receipt in receipts:
            if 'file_path' not in receipt:
                continue

            source = Path(receipt['file_path'])

            # Если путь относительный - делаем абсолютным
            if not source.is_absolute():
                from ..config import settings
                source = settings.BASE_DIR / source

            if source.exists():
                dest = receipts_folder / f"{receipt.get('category', 'other')}_{source.name}"
                shutil.copy(source, dest)

    def _create_zip(self, trip_folder: Path) -> str:
        """Создает ZIP архив ТОЛЬКО с чеками (не документами)"""
        zip_path = trip_folder / "receipts.zip"
        receipts_folder = trip_folder / "receipts"

        # Удаляем старый ZIP
        if zip_path.exists():
            zip_path.unlink()

        # Архивируем ТОЛЬКО папку receipts
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            if receipts_folder.exists():
                for file_path in receipts_folder.rglob('*'):
                    if file_path.is_file():
                        arcname = file_path.relative_to(receipts_folder)
                        zipf.write(file_path, arcname)

        return str(zip_path)


def calculate_per_diem_days(date_from, date_to, departure_time, arrival_time) -> float:
    """Расчет суточных с коэффициентами"""

    # Коэффициент выезда
    if departure_time:
        if departure_time.hour < 12:
            departure_coef = 1.0
        elif 12 <= departure_time.hour < 18:
            departure_coef = 0.5
        else:
            departure_coef = 0.4
    else:
        departure_coef = 1.0

    # Коэффициент приезда
    if arrival_time:
        if arrival_time.hour >= 18:
            arrival_coef = 1.0
        elif 12 <= arrival_time.hour < 18:
            arrival_coef = 0.5
        else:
            arrival_coef = 0.4
    else:
        arrival_coef = 1.0

    # Полные дни
    days_diff = (date_to - date_from).days
    full_days = max(0, days_diff - 1)

    total_days = departure_coef + full_days + arrival_coef

    return total_days
