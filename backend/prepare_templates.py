"""
Скрипт для подготовки реальных шаблонов документов
- Конвертация .xls -> .xlsx
- Конвертация .doc -> .docx
- Переименование в правильный формат
- Анализ структуры документов
"""
import os
from pathlib import Path
import shutil

# Импорты для конвертации
try:
    from win32com import client as win32
    WORD_AVAILABLE = True
    EXCEL_AVAILABLE = True
except ImportError:
    print("! win32com не установлен. Конвертация .doc/.xls невозможна")
    print("  Установите: pip install pywin32")
    WORD_AVAILABLE = False
    EXCEL_AVAILABLE = False

from openpyxl import load_workbook
from docx import Document


def convert_xls_to_xlsx(xls_path, xlsx_path):
    """Конвертация .xls -> .xlsx"""
    if not EXCEL_AVAILABLE:
        print(f"! Пропуск конвертации {xls_path} - win32com не установлен")
        return False

    try:
        excel = win32.Dispatch("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False

        # Открываем .xls файл
        wb = excel.Workbooks.Open(str(xls_path.absolute()))

        # Сохраняем как .xlsx (формат 51)
        wb.SaveAs(str(xlsx_path.absolute()), FileFormat=51)
        wb.Close()
        excel.Quit()

        print(f"[OK] Конвертировано: {xls_path.name} -> {xlsx_path.name}")
        return True
    except Exception as e:
        print(f"[ERR] Ошибка конвертации {xls_path.name}: {e}")
        return False


def convert_doc_to_docx(doc_path, docx_path):
    """Конвертация .doc -> .docx"""
    if not WORD_AVAILABLE:
        print(f"! Пропуск конвертации {doc_path} - win32com не установлен")
        return False

    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False

        # Открываем .doc файл
        doc = word.Documents.Open(str(doc_path.absolute()))

        # Сохраняем как .docx (формат 16)
        doc.SaveAs(str(docx_path.absolute()), FileFormat=16)
        doc.Close()
        word.Quit()

        print(f"[OK] Конвертировано: {doc_path.name} -> {docx_path.name}")
        return True
    except Exception as e:
        print(f"[ERR] Ошибка конвертации {doc_path.name}: {e}")
        return False


def analyze_excel_template(xlsx_path):
    """Анализ структуры Excel шаблона"""
    print(f"\n=== Анализ: {xlsx_path.name} ===")

    try:
        wb = load_workbook(xlsx_path, data_only=False)
        ws = wb.active

        print(f"  Листов: {len(wb.sheetnames)}")
        print(f"  Активный лист: {ws.title}")
        print(f"  Размер: {ws.max_row} строк × {ws.max_column} колонок")

        # Named Ranges
        if hasattr(wb, 'defined_names'):
            names = list(wb.defined_names.definedName)
            if names:
                print(f"  Named Ranges ({len(names)}):")
                for name in names[:10]:  # Первые 10
                    print(f"    - {name.name}")

        # Находим заполненные ячейки
        print(f"  Примеры данных:")
        sample_count = 0
        for row in ws.iter_rows(max_row=50, max_col=20):
            for cell in row:
                if cell.value and isinstance(cell.value, str) and len(cell.value) > 3:
                    print(f"    {cell.coordinate}: {str(cell.value)[:50]}")
                    sample_count += 1
                    if sample_count >= 10:
                        break
            if sample_count >= 10:
                break

        wb.close()
        return True
    except Exception as e:
        print(f"  [ERR] Ошибка анализа: {e}")
        return False


def analyze_word_template(docx_path):
    """Анализ структуры Word шаблона"""
    print(f"\n=== Анализ: {docx_path.name} ===")

    try:
        doc = Document(docx_path)

        print(f"  Параграфов: {len(doc.paragraphs)}")
        print(f"  Таблиц: {len(doc.tables)}")

        # Примеры текста
        print(f"  Примеры текста:")
        for i, para in enumerate(doc.paragraphs[:15]):
            text = para.text.strip()
            if text:
                print(f"    [{i+1}] {text[:60]}")

        return True
    except Exception as e:
        print(f"  [ERR] Ошибка анализа: {e}")
        return False


def main():
    """Основная функция подготовки шаблонов"""
    print("=" * 60)
    print("ПОДГОТОВКА РЕАЛЬНЫХ ШАБЛОНОВ ДОКУМЕНТОВ")
    print("=" * 60)

    templates_dir = Path(__file__).parent / "templates"

    # 1. Удаляем старые тестовые шаблоны
    print("\n1. Удаление тестовых шаблонов...")
    test_files = [
        "ao_template.xlsx",
        "prikaz_template.docx",
        "sutochnye_template.xlsx",
        "sz_template.docx"
    ]
    for f in test_files:
        file_path = templates_dir / f
        if file_path.exists():
            file_path.unlink()
            print(f"  [OK] Удален: {f}")

    # 2. Конвертация файлов
    print("\n2. Конвертация документов...")

    # АО .xls -> .xlsx
    ao_xls = templates_dir / "АО 20.06.2025 — копия.xls"
    ao_xlsx = templates_dir / "ao_template.xlsx"
    if ao_xls.exists():
        convert_xls_to_xlsx(ao_xls, ao_xlsx)

    # Суточные .xls -> .xlsx
    sutochnie_xls = templates_dir / "Суточные 20.06.25 — копия.xls"
    sutochnie_xlsx = templates_dir / "sutochnye_template.xlsx"
    if sutochnie_xls.exists():
        convert_xls_to_xlsx(sutochnie_xls, sutochnie_xlsx)

    # СЗ .doc -> .docx
    sz_doc = templates_dir / "СЗ 20.06.2025 — копия.doc"
    sz_docx = templates_dir / "sz_template.docx"
    if sz_doc.exists():
        convert_doc_to_docx(sz_doc, sz_docx)

    # Приказ уже .docx - просто копируем
    prikaz_source = templates_dir / "Приказ 13.06.2025 — копия.docx"
    prikaz_dest = templates_dir / "prikaz_template.docx"
    if prikaz_source.exists():
        shutil.copy(prikaz_source, prikaz_dest)
        print(f"[OK] Скопирован: {prikaz_source.name} -> {prikaz_dest.name}")

    # 3. Анализ структуры
    print("\n3. Анализ структуры шаблонов...")

    if ao_xlsx.exists():
        analyze_excel_template(ao_xlsx)

    if sutochnie_xlsx.exists():
        analyze_excel_template(sutochnie_xlsx)

    if prikaz_dest.exists():
        analyze_word_template(prikaz_dest)

    if sz_docx.exists():
        analyze_word_template(sz_docx)

    print("\n" + "=" * 60)
    print("ГОТОВО!")
    print("=" * 60)
    print("\nШаблоны подготовлены:")
    print("  [OK] ao_template.xlsx - Авансовый отчет")
    print("  [OK] sutochnye_template.xlsx - Суточные")
    print("  [OK] prikaz_template.docx - Приказ")
    print("  [OK] sz_template.docx - Служебная записка")
    print("\nСледующие шаги:")
    print("  1. Проверьте что все файлы корректно конвертированы")
    print("  2. Настройте генератор документов под новые шаблоны")
    print("  3. Добавьте маркеры {{переменная}} в нужные места")


if __name__ == "__main__":
    main()
